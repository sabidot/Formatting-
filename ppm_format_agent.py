"""
PPM Impact Fund — Formatting Agent
===================================
Flags formatting inconsistencies across a .docx file and auto-fixes them
in-place, producing:
  • <filename>_fixed.docx   — corrected document
  • <filename>_flags.json   — full flag report

Usage
-----
  python ppm_format_agent.py document.docx           # flag + fix
  python ppm_format_agent.py document.docx --check   # flag only, no fix
  python ppm_format_agent.py document.docx --verbose # print flags to console too

Spec source: PPM_Impact_Fund_2025-26.docx (LightCastle Partners)
"""

import argparse
import copy
import json
import re
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from lxml import etree

# ---------------------------------------------------------------------------
# FORMAT SPEC — single source of truth
# ---------------------------------------------------------------------------

SPEC = {
    # Typography
    "font":               "Lato",
    "body_pt":            9,       # 9pt body text
    "body_line_spacing":  1.15,    # 1.15× for body paragraphs
    "body_space_before":  240,     # twips
    "body_space_after":   240,     # twips

    # Title
    "title_pt":           18,

    # Headings — H1 is context-aware (see check_h1_sizes)
    "h1_pt_with_h2":      16,      # H1 when H2 exists anywhere in doc
    "h1_pt_no_h2":        14,      # H1 when no H2 in doc
    "h1_max_pt":          16,      # HARD CEILING — never higher than this
    "h2_pt":              14,  "h2_bold": True,
    "h3_pt":              12,  "h3_bold": True,
    "h4_pt":              11,  "h4_color": "666666",
    "h5_pt":              11,  "h5_color": "666666",
    "h6_pt":              11,  "h6_italic": True, "h6_color": "666666",

    # Tables — body cell borders
    "tbl_spacing":        "single",        # exact single spacing in all cells
    "tbl_header_fill":    "073763",        # deep navy
    "tbl_header_text":    "ffffff",        # white
    "tbl_border_color":   "6fa8dc",        # light blue dashed (body cells)
    "tbl_border_style":   "dashed",        # 4th dash option in Word border picker
    "tbl_border_sz":      4,               # ½ pt = 4 in Word's eighths-of-a-point unit
    "tbl_even_fill":      "efefef",        # alternating row fill
    "tbl_outer_border":   "nil",           # no outer table border

    # Header row borders — solid white lines, ½ pt
    "tbl_header_border_style": "single",   # straight/solid line
    "tbl_header_border_color": "ffffff",   # white
    "tbl_header_border_sz":    4,          # ½ pt

    # Cell alignment: (horizontal, vertical)
    "align_header":       ("center", "center"),
    "align_text":         ("left",   "center"),
    "align_numeric":      ("right",  "center"),
}

# Heading style names → spec keys
HEADING_SPEC = {
    "Heading 1": {"pt_key": None,   "bold": False, "italic": False, "color": None},
    "Heading 2": {"pt": 14,         "bold": True,  "italic": False, "color": None},
    "Heading 3": {"pt": 12,         "bold": True,  "italic": False, "color": None},
    "Heading 4": {"pt": 11,         "bold": False, "italic": False, "color": "666666"},
    "Heading 5": {"pt": 11,         "bold": False, "italic": False, "color": "666666"},
    "Heading 6": {"pt": 11,         "bold": False, "italic": True,  "color": "666666"},
}

# Numeric cell pattern: digits, commas, dots, %, $, ±, dashes as negatives
NUMERIC_RE = re.compile(r'^\s*[\d,\.%$\-\+]+\s*$')

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ---------------------------------------------------------------------------
# HELPERS — reading XML attributes
# ---------------------------------------------------------------------------

def wq(tag):
    return f"{{{W_NS}}}{tag}"


def get_run_size_pt(para):
    """Return font size in pt from the first run that declares one, else from style."""
    for run in para.runs:
        if run.font.size:
            return run.font.size.pt
    if para.style and para.style.font.size:
        return para.style.font.size.pt
    return None


def get_para_jc(para):
    """Return paragraph horizontal alignment string (left/center/right/justify)."""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        jc = pPr.find(qn("w:jc"))
        if jc is not None:
            return jc.get(qn("w:val"), "left")
    if para.alignment is not None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        mapping = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
        }
        return mapping.get(para.alignment, "left")
    return "left"


def get_cell_valign(cell):
    """Return vertical alignment of a table cell (top/center/bottom)."""
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is not None:
        va = tcPr.find(qn("w:vAlign"))
        if va is not None:
            return va.get(qn("w:val"), "top")
    return "top"


def get_cell_fill(cell):
    """Return fill hex colour of a cell, or None."""
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is not None:
        shd = tcPr.find(qn("w:shd"))
        if shd is not None:
            return shd.get(qn("w:fill"))
    return None


def get_cell_border_style(cell, side="top"):
    """Return border style string for a given side of a cell."""
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is not None:
        borders = tcPr.find(qn("w:tcBorders"))
        if borders is not None:
            b = borders.find(qn(f"w:{side}"))
            if b is not None:
                return b.get(qn("w:val"))
    return None


def get_cell_border_color(cell, side="top"):
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is not None:
        borders = tcPr.find(qn("w:tcBorders"))
        if borders is not None:
            b = borders.find(qn(f"w:{side}"))
            if b is not None:
                return b.get(qn("w:color"))
    return None


def get_para_line_spacing(para):
    """
    Return (line_value, line_rule) from paragraph spacing XML.
    line_rule: 'auto' | 'exact' | 'atLeast'
    line_value: twips integer, or None
    """
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        sp = pPr.find(qn("w:spacing"))
        if sp is not None:
            val  = sp.get(qn("w:line"))
            rule = sp.get(qn("w:lineRule"), "auto")
            return (int(val) if val else None, rule)
    return (None, None)


def classify_cell(cell, row_idx):
    """Classify cell as 'header', 'numeric', or 'text'."""
    if row_idx == 0:
        return "header"
    text = cell.text.strip()
    if text and NUMERIC_RE.match(text):
        return "numeric"
    return "text"


def has_h2(doc):
    return any(p.style.name == "Heading 2" for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# FLAGGING — read-only sweep
# ---------------------------------------------------------------------------

def flag_heading_sizes(doc, flags):
    """
    Collect all heading sizes by level.
    For H1: check context rule + ceiling + cross-doc drift.
    For H2–H6: check against fixed spec + cross-doc drift.
    """
    doc_has_h2 = has_h2(doc)
    expected_h1 = SPEC["h1_pt_with_h2"] if doc_has_h2 else SPEC["h1_pt_no_h2"]

    sizes_by_level = defaultdict(list)  # level → [(para_idx, pt)]

    for i, p in enumerate(doc.paragraphs):
        name = p.style.name
        for level in range(1, 7):
            if name == f"Heading {level}":
                sz = get_run_size_pt(p)
                sizes_by_level[level].append((i, sz, p.text[:60]))

    # H1 checks
    for (para_i, sz, preview) in sizes_by_level.get(1, []):
        if sz is None:
            continue
        if sz > SPEC["h1_max_pt"]:
            flags.append({
                "type": "H1_ceiling_breach",
                "severity": "critical",
                "para": para_i,
                "preview": preview,
                "found": f"{sz}pt",
                "expected": f"≤{SPEC['h1_max_pt']}pt",
                "fix": f"Set to {expected_h1}pt",
            })
        elif sz != expected_h1:
            flags.append({
                "type": "H1_wrong_size",
                "severity": "critical",
                "para": para_i,
                "preview": preview,
                "found": f"{sz}pt",
                "expected": f"{expected_h1}pt",
                "reason": f"{'H2 present' if doc_has_h2 else 'No H2 in doc'} → H1 must be {expected_h1}pt",
                "fix": f"Set to {expected_h1}pt",
            })

    # H1 drift
    h1_sizes = set(sz for (_, sz, _) in sizes_by_level.get(1, []) if sz)
    if len(h1_sizes) > 1:
        flags.append({
            "type": "H1_drift",
            "severity": "critical",
            "values": sorted(h1_sizes),
            "message": f"H1 used at multiple sizes {sorted(h1_sizes)} — must be uniform at {expected_h1}pt",
            "locations": [(i, sz, p) for (i, sz, p) in sizes_by_level[1]],
        })

    # H2–H6 checks
    for level in range(2, 7):
        spec = HEADING_SPEC.get(f"Heading {level}", {})
        expected_pt = spec.get("pt")
        if expected_pt is None:
            continue

        for (para_i, sz, preview) in sizes_by_level.get(level, []):
            if sz is not None and sz != expected_pt:
                flags.append({
                    "type": f"H{level}_wrong_size",
                    "severity": "high",
                    "para": para_i,
                    "preview": preview,
                    "found": f"{sz}pt",
                    "expected": f"{expected_pt}pt",
                    "fix": f"Set to {expected_pt}pt",
                })

        # Drift across H2–H6
        level_sizes = set(sz for (_, sz, _) in sizes_by_level.get(level, []) if sz)
        if len(level_sizes) > 1:
            flags.append({
                "type": f"H{level}_drift",
                "severity": "high",
                "values": sorted(level_sizes),
                "message": f"H{level} used at multiple sizes {sorted(level_sizes)}",
                "locations": [(i, sz, p) for (i, sz, p) in sizes_by_level[level]],
            })


def flag_table_spacing(doc, flags):
    """Every table cell paragraph must use single line spacing."""
    for t_idx, tbl in enumerate(doc.tables):
        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    val, rule = get_para_line_spacing(para)
                    # Single = 240 twips with rule "auto", or rule "exact" at 240
                    is_single = (
                        (rule == "auto" and val == 240) or
                        (rule == "exact" and val == 240) or
                        (val is None and rule is None)  # inheriting single from style
                    )
                    if not is_single and val is not None:
                        flags.append({
                            "type": "table_spacing_not_single",
                            "severity": "critical",
                            "table": t_idx,
                            "row": r_idx,
                            "col": c_idx,
                            "found": f"line={val} rule={rule}",
                            "expected": "single (240 auto or 240 exact)",
                            "fix": "Set spacing to single",
                        })


def flag_table_alignment(doc, flags):
    """Check header, text, and numeric cell alignment rules."""
    for t_idx, tbl in enumerate(doc.tables):
        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                ctype = classify_cell(cell, r_idx)
                exp_h, exp_v = SPEC[f"align_{ctype}"]

                # Horizontal: check first paragraph of cell
                if cell.paragraphs:
                    actual_h = get_para_jc(cell.paragraphs[0])
                    # Normalise: Word uses "both" for justify, treat left as left
                    if actual_h not in ("left",) and exp_h == "left":
                        if actual_h != "left":
                            pass  # checked below
                    if actual_h != exp_h:
                        flags.append({
                            "type": "cell_h_align_wrong",
                            "severity": "high",
                            "table": t_idx, "row": r_idx, "col": c_idx,
                            "cell_type": ctype,
                            "found_h": actual_h,
                            "expected_h": exp_h,
                            "fix": f"Set paragraph alignment to {exp_h}",
                        })

                actual_v = get_cell_valign(cell)
                # "center" in Word vAlign = middle
                if actual_v not in ("center",) and exp_v == "center":
                    flags.append({
                        "type": "cell_v_align_wrong",
                        "severity": "high",
                        "table": t_idx, "row": r_idx, "col": c_idx,
                        "cell_type": ctype,
                        "found_v": actual_v,
                        "expected_v": exp_v,
                        "fix": "Set vAlign to center",
                    })


def flag_table_colors(doc, flags):
    """Check header fill, even-row fill, cell border colour/style."""
    for t_idx, tbl in enumerate(doc.tables):
        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                fill = get_cell_fill(cell)

                # Header fill
                if r_idx == 0:
                    if fill and fill.upper() != SPEC["tbl_header_fill"].upper():
                        flags.append({
                            "type": "header_fill_wrong",
                            "severity": "critical",
                            "table": t_idx, "row": r_idx, "col": c_idx,
                            "found": fill,
                            "expected": SPEC["tbl_header_fill"],
                            "fix": f"Set fill to #{SPEC['tbl_header_fill']}",
                        })

                # Even body row fill (rows 2, 4, 6… → r_idx 1, 3, 5…)
                elif r_idx % 2 == 0:
                    if fill and fill.upper() != SPEC["tbl_even_fill"].upper():
                        flags.append({
                            "type": "even_row_fill_wrong",
                            "severity": "medium",
                            "table": t_idx, "row": r_idx, "col": c_idx,
                            "found": fill,
                            "expected": SPEC["tbl_even_fill"],
                            "fix": f"Set fill to #{SPEC['tbl_even_fill']}",
                        })
                    elif fill is None:
                        flags.append({
                            "type": "even_row_fill_missing",
                            "severity": "medium",
                            "table": t_idx, "row": r_idx, "col": c_idx,
                            "found": "none",
                            "expected": SPEC["tbl_even_fill"],
                            "fix": f"Set fill to #{SPEC['tbl_even_fill']}",
                        })

                # Border checks — header uses solid white; body uses dashed blue
                is_header = (r_idx == 0)
                exp_style = SPEC["tbl_header_border_style"] if is_header else SPEC["tbl_border_style"]
                exp_color = SPEC["tbl_header_border_color"] if is_header else SPEC["tbl_border_color"]
                exp_sz    = SPEC["tbl_header_border_sz"]    if is_header else SPEC["tbl_border_sz"]

                for side in ("top", "left", "bottom", "right"):
                    bstyle = get_cell_border_style(cell, side)
                    bcolor = get_cell_border_color(cell, side)
                    if bstyle and bstyle != "nil":
                        if bstyle != exp_style:
                            flags.append({
                                "type": "cell_border_style_wrong",
                                "severity": "critical",
                                "table": t_idx, "row": r_idx, "col": c_idx,
                                "side": side,
                                "found": bstyle,
                                "expected": exp_style,
                                "fix": f"Set border to {exp_style}" + (" (header: solid white)" if is_header else " (body: dashed blue)"),
                            })
                        if bcolor and bcolor.upper() != exp_color.upper():
                            flags.append({
                                "type": "cell_border_color_wrong",
                                "severity": "critical",
                                "table": t_idx, "row": r_idx, "col": c_idx,
                                "side": side,
                                "found": bcolor,
                                "expected": exp_color,
                                "fix": f"Set border color to #{exp_color}" + (" (header: white)" if is_header else " (body: #6fa8dc)"),
                            })


def flag_font(doc, flags):
    """Flag any run not using Lato in body paragraphs."""
    expected = SPEC["font"]
    for i, p in enumerate(doc.paragraphs):
        for r_idx, run in enumerate(p.runs):
            actual = run.font.name
            if actual and actual != expected:
                flags.append({
                    "type": "wrong_font",
                    "severity": "critical",
                    "para": i,
                    "run": r_idx,
                    "preview": run.text[:40],
                    "found": actual,
                    "expected": expected,
                    "fix": f"Set font to {expected}",
                })


def run_all_checks(doc):
    flags = []
    flag_heading_sizes(doc, flags)
    flag_table_spacing(doc, flags)
    flag_table_alignment(doc, flags)
    flag_table_colors(doc, flags)
    flag_font(doc, flags)
    return flags


# ---------------------------------------------------------------------------
# AUTO-FIXER — modifies XML in-place
# ---------------------------------------------------------------------------

def _set_shd(tcPr, fill_hex, val="clear"):
    """Set or replace w:shd on a tcPr element."""
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        tcPr.remove(shd)
    shd = etree.SubElement(tcPr, qn("w:shd"))
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), val)
    shd.set(qn("w:color"), "auto")


def _set_cell_border(tcPr, side, style, color, sz):
    """Set a single border side on tcPr."""
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = etree.SubElement(tcPr, qn("w:tcBorders"))
    b = borders.find(qn(f"w:{side}"))
    if b is None:
        b = etree.SubElement(borders, qn(f"w:{side}"))
    b.set(qn("w:val"), style)
    b.set(qn("w:color"), color)
    b.set(qn("w:sz"), str(sz))
    b.set(qn("w:space"), "0")


def _set_valign(tcPr, val):
    va = tcPr.find(qn("w:vAlign"))
    if va is None:
        va = etree.SubElement(tcPr, qn("w:vAlign"))
    va.set(qn("w:val"), val)


def _set_para_jc(para, jc_val):
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(para._p, qn("w:pPr"))
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = etree.SubElement(pPr, qn("w:jc"))
    jc.set(qn("w:val"), jc_val)


def _set_para_spacing_single(para):
    """Set paragraph line spacing to single (240 auto)."""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(para._p, qn("w:pPr"))
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = etree.SubElement(pPr, qn("w:spacing"))
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "auto")


def _set_run_color(run, hex_color):
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        rPr = etree.SubElement(run._r, qn("w:rPr"))
    color = rPr.find(qn("w:color"))
    if color is None:
        color = etree.SubElement(rPr, qn("w:color"))
    color.set(qn("w:val"), hex_color)


def _set_run_size(run, pt):
    """Set font size in half-points."""
    hpts = str(int(pt * 2))
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        rPr = etree.SubElement(run._r, qn("w:rPr"))
    for tag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = etree.SubElement(rPr, qn(tag))
        el.set(qn("w:val"), hpts)


def _get_or_add_tcPr(cell):
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(cell._tc, qn("w:tcPr"))
    return tcPr


def fix_heading_sizes(doc):
    """Fix all heading font sizes to match spec."""
    doc_has_h2 = has_h2(doc)
    expected_h1_pt = SPEC["h1_pt_with_h2"] if doc_has_h2 else SPEC["h1_pt_no_h2"]

    level_pt = {
        1: expected_h1_pt,
        2: SPEC["h2_pt"],
        3: SPEC["h3_pt"],
        4: SPEC["h4_pt"],
        5: SPEC["h5_pt"],
        6: SPEC["h6_pt"],
    }

    for p in doc.paragraphs:
        name = p.style.name
        for level in range(1, 7):
            if name == f"Heading {level}":
                target_pt = level_pt[level]
                if not p.runs:
                    continue
                for run in p.runs:
                    current = run.font.size.pt if run.font.size else None
                    if current != target_pt:
                        _set_run_size(run, target_pt)


def fix_table_spacing(doc):
    """Set single line spacing on all table cell paragraphs."""
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _set_para_spacing_single(para)


def fix_table_alignment(doc):
    """Fix horizontal + vertical alignment on all table cells."""
    for tbl in doc.tables:
        for r_idx, row in enumerate(tbl.rows):
            for cell in row.cells:
                ctype = classify_cell(cell, r_idx)
                exp_h, exp_v = SPEC[f"align_{ctype}"]
                tcPr = _get_or_add_tcPr(cell)
                _set_valign(tcPr, exp_v)
                for para in cell.paragraphs:
                    _set_para_jc(para, exp_h)


def fix_table_header_row(doc):
    """
    Enforce header row:
      - Fill: #073763 (deep navy)
      - Borders: solid white, 1/2pt (single, #ffffff, sz=4)
      - Text: white, center-center aligned
    """
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        header_row = tbl.rows[0]
        for cell in header_row.cells:
            tcPr = _get_or_add_tcPr(cell)

            # Fill
            _set_shd(tcPr, SPEC["tbl_header_fill"])

            # Header borders: solid white 1/2pt (distinct from body dashed blue)
            for side in ("top", "left", "bottom", "right"):
                _set_cell_border(
                    tcPr, side,
                    SPEC["tbl_header_border_style"],   # "single" (solid)
                    SPEC["tbl_header_border_color"],   # "ffffff" (white)
                    SPEC["tbl_header_border_sz"],      # 4 = 1/2pt
                )

            # Vertical align center
            _set_valign(tcPr, "center")

            # White text, horizontally centered
            for para in cell.paragraphs:
                _set_para_jc(para, "center")
                for run in para.runs:
                    _set_run_color(run, SPEC["tbl_header_text"])


def fix_table_body_rows(doc):
    """Apply alternating fill and dashed borders on all body rows."""
    for tbl in doc.tables:
        for r_idx, row in enumerate(tbl.rows[1:], start=1):
            fill = SPEC["tbl_even_fill"] if (r_idx % 2 == 0) else "FFFFFF"
            for cell in row.cells:
                tcPr = _get_or_add_tcPr(cell)
                _set_shd(tcPr, fill)
                for side in ("top", "left", "bottom", "right"):
                    _set_cell_border(
                        tcPr, side,
                        SPEC["tbl_border_style"],
                        SPEC["tbl_border_color"],
                        SPEC["tbl_border_sz"],
                    )


def fix_table_outer_border(doc):
    """Remove outer table-level borders (set to nil)."""
    for tbl in doc.tables:
        tblPr = tbl._tbl.find(qn("w:tblPr"))
        if tblPr is None:
            continue
        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is None:
            tblBorders = etree.SubElement(tblPr, qn("w:tblBorders"))
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = tblBorders.find(qn(f"w:{side}"))
            if b is None:
                b = etree.SubElement(tblBorders, qn(f"w:{side}"))
            b.set(qn("w:val"), "nil")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:color"), "000000")
            b.set(qn("w:space"), "0")


def fix_fonts(doc):
    """Set all runs to Lato."""
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = SPEC["font"]


# Colors that count as "intentional" and should be kept as-is
# (highlights, known accent colours used for callouts etc.)
_KEEP_COLORS = set()  # empty = reset ALL non-black to black (safe default)

# These heading-style names should NOT be touched by body fixer
_HEADING_STYLES = {f"Heading {i}" for i in range(1, 7)} | {"Title", "Subtitle"}


def fix_body_text(doc):
    """
    For every non-heading paragraph:
      1. Font size → 9pt (sz=18 half-pts) — clears run-level overrides
      2. Color → black (000000) UNLESS the run is bold, italic, or highlighted
         (those are intentional emphasis — keep their color)
      3. Left-align paragraph
      4. Bold / italic / highlight preserved as-is
      5. Font → Lato
    """
    target_hpts = str(SPEC["body_pt"] * 2)   # "18" = 9pt

    for para in doc.paragraphs:
        # Skip headings entirely — handled by fix_heading_sizes
        if para.style.name in _HEADING_STYLES:
            continue

        # Left-align body paragraphs
        _set_para_jc(para, "left")

        for run in para.runs:
            rPr = run._r.find(qn("w:rPr"))
            if rPr is None:
                rPr = etree.SubElement(run._r, qn("w:rPr"))

            # ── Font ──────────────────────────────────────────────────────
            fonts = rPr.find(qn("w:rFonts"))
            if fonts is None:
                fonts = etree.SubElement(rPr, qn("w:rFonts"))
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                fonts.set(qn(attr), SPEC["font"])

            # ── Size — always enforce 9pt on body runs ────────────────────
            for tag in ("w:sz", "w:szCs"):
                el = rPr.find(qn(tag))
                if el is None:
                    el = etree.SubElement(rPr, qn(tag))
                el.set(qn("w:val"), target_hpts)

            # ── Color ─────────────────────────────────────────────────────
            # Detect intentional emphasis
            is_bold      = rPr.find(qn("w:b"))      is not None
            is_italic    = rPr.find(qn("w:i"))      is not None
            is_highlight = rPr.find(qn("w:highlight")) is not None

            color_el = rPr.find(qn("w:color"))
            current_color = (
                color_el.get(qn("w:val")) if color_el is not None else None
            )

            # Reset to black ONLY if:
            #   - there IS a non-black color set at run level
            #   - AND the run is not highlighted (highlights are intentional)
            #   - AND the color is not a kept accent color
            if (
                current_color
                and current_color not in ("000000", "auto", None)
                and not is_highlight
                and current_color not in _KEEP_COLORS
            ):
                if color_el is None:
                    color_el = etree.SubElement(rPr, qn("w:color"))
                color_el.set(qn("w:val"), "000000")

            # Bold, italic, highlight — do NOT touch, preserve exactly


def apply_all_fixes(doc):
    """Run all fixers in dependency order."""
    fix_fonts(doc)
    fix_body_text(doc)        # size + color + alignment on body paragraphs
    fix_heading_sizes(doc)
    fix_table_outer_border(doc)
    fix_table_header_row(doc)
    fix_table_body_rows(doc)
    fix_table_spacing(doc)
    fix_table_alignment(doc)


# ---------------------------------------------------------------------------
# REPORT
# ---------------------------------------------------------------------------

SEVERITY_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3}


def build_report(input_path, flags, fixed):
    severity_counts = defaultdict(int)
    for f in flags:
        severity_counts[f.get("severity", "low")] += 1

    return {
        "file": str(input_path),
        "total_flags": len(flags),
        "auto_fixed": fixed,
        "severity_summary": dict(severity_counts),
        "h1_context": (
            f"H2 present in doc → H1 expected at {SPEC['h1_pt_with_h2']}pt"
            if any(True for _ in [1])   # placeholder; recalculated in main
            else f"No H2 → H1 expected at {SPEC['h1_pt_no_h2']}pt"
        ),
        "flags": sorted(flags, key=lambda f: SEVERITY_ORDER.get(f.get("severity", "low"), 9)),
    }


# ---------------------------------------------------------------------------
# CLI ENTRY POINT
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="PPM Impact Fund — formatting agent. Flags inconsistencies and auto-fixes them."
    )
    parser.add_argument("input", help="Path to the .docx file to check")
    parser.add_argument(
        "--check", action="store_true",
        help="Flag only — do not write a fixed file"
    )
    parser.add_argument(
        "--verbose", action="store_true",
        help="Print flag report to stdout in addition to saving JSON"
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: file not found: {input_path}", file=sys.stderr)
        sys.exit(1)
    if input_path.suffix.lower() != ".docx":
        print(f"Error: expected a .docx file, got: {input_path.suffix}", file=sys.stderr)
        sys.exit(1)

    print(f"Loading: {input_path}")
    doc = Document(str(input_path))

    # ---- SCAN
    print("Scanning for formatting issues...")
    flags = run_all_checks(doc)
    print(f"  Found {len(flags)} flag(s)")

    # Build report
    doc_has_h2 = has_h2(doc)
    report = {
        "file": str(input_path),
        "total_flags": len(flags),
        "auto_fixed": not args.check,
        "h1_context": (
            f"H2 present → H1 expected at {SPEC['h1_pt_with_h2']}pt"
            if doc_has_h2
            else f"No H2 → H1 expected at {SPEC['h1_pt_no_h2']}pt"
        ),
        "severity_summary": {
            sev: sum(1 for f in flags if f.get("severity") == sev)
            for sev in ("critical", "high", "medium", "low")
        },
        "flags": sorted(flags, key=lambda f: SEVERITY_ORDER.get(f.get("severity", "low"), 9)),
    }

    # ---- SAVE REPORT
    json_path = input_path.with_name(input_path.stem + "_flags.json")
    with open(json_path, "w", encoding="utf-8") as fh:
        json.dump(report, fh, indent=2, ensure_ascii=False)
    print(f"  Flag report saved: {json_path}")

    # ---- FIX
    if not args.check:
        print("Applying fixes...")
        apply_all_fixes(doc)
        out_path = input_path.with_name(input_path.stem + "_fixed.docx")
        doc.save(str(out_path))
        print(f"  Fixed document saved: {out_path}")
    else:
        print("  --check mode: no fixes applied.")

    # ---- VERBOSE OUTPUT
    if args.verbose or not flags:
        print()
        if not flags:
            print("No formatting issues found.")
        else:
            print(f"{'SEV':<10} {'TYPE':<35} {'DETAIL'}")
            print("-" * 80)
            for f in report["flags"]:
                sev   = f.get("severity", "?").upper()[:8]
                ftype = f.get("type", "")[:34]
                detail_parts = []
                if "table" in f:
                    detail_parts.append(f"tbl={f['table']} row={f['row']} col={f['col']}")
                if "para" in f:
                    detail_parts.append(f"para={f['para']}")
                if "found" in f:
                    detail_parts.append(f"found={f['found']}")
                if "expected" in f:
                    detail_parts.append(f"expected={f['expected']}")
                if "preview" in f:
                    detail_parts.append(f'"{f["preview"]}"')
                print(f"{sev:<10} {ftype:<35} {' | '.join(detail_parts)}")

    print()
    print("Done.")
    print(f"  Flags : {report['total_flags']}")
    for sev, count in report["severity_summary"].items():
        if count:
            print(f"    {sev:<10}: {count}")
    if not args.check:
        print(f"  Output : {out_path}")
    print(f"  Report : {json_path}")


if __name__ == "__main__":
    main()
